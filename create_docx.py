from docx import Document

def main():
    f = open("text.txt", "r", encoding="utf8")
    content = f.read()
    document = Document()
    # document.add_heading('A simple text', level=1)
    document.add_paragraph(content)
    document.save('docx_file.docx')
if __name__ == "__main__":
    main() 
